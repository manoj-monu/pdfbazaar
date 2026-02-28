import os
from io import BytesIO
from typing import List, Optional
from fastapi import FastAPI, File, UploadFile, Form, HTTPException, Response
from fastapi.middleware.cors import CORSMiddleware
from pypdf import PdfReader, PdfWriter
import fitz  # PyMuPDF
from PIL import Image
import zipfile
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import io

app = FastAPI(title="iLovePDF Clone API")

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
def read_root():
    return {"status": "ok", "message": "iLovePDF clone backend"}

@app.post("/api/merge")
async def merge_pdfs(files: List[UploadFile] = File(...)):
    if len(files) < 2:
        raise HTTPException(status_code=400, detail="Please upload at least two PDFs to merge.")
    merger = PdfWriter()
    for file in files:
        if not file.filename.lower().endswith('.pdf'):
            raise HTTPException(status_code=400, detail=f"File {file.filename} is not a PDF")
        content = await file.read()
        merger.append(BytesIO(content))
    output = BytesIO()
    merger.write(output)
    merger.close()
    output.seek(0)
    return Response(content=output.read(), media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=merged.pdf"})

@app.post("/api/split")
async def split_pdf(files: List[UploadFile] = File(...)):
    if len(files) != 1:
        raise HTTPException(status_code=400, detail="Please upload exactly one PDF.")
    content = await files[0].read()
    reader = PdfReader(BytesIO(content))
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
        for i in range(len(reader.pages)):
            writer = PdfWriter()
            writer.add_page(reader.pages[i])
            page_buffer = BytesIO()
            writer.write(page_buffer)
            zip_file.writestr(f"page_{i + 1}.pdf", page_buffer.getvalue())
    zip_buffer.seek(0)
    return Response(content=zip_buffer.read(), media_type="application/zip", headers={"Content-Disposition": "attachment; filename=split_pages.zip"})

@app.post("/api/compress")
async def compress_pdf(files: List[UploadFile] = File(...)):
    if len(files) != 1:
        raise HTTPException(status_code=400, detail="Please upload exactly one PDF.")
    content = await files[0].read()
    try:
        doc = fitz.open("pdf", content)
        out_bytes = doc.write(garbage=4, deflate=True)
        doc.close()
    except Exception as e:
        raise HTTPException(status_code=500, detail="Compression failed: " + str(e))
    return Response(content=out_bytes, media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=compressed.pdf"})

@app.post("/api/pdf-to-jpg")
async def pdf_to_jpg(files: List[UploadFile] = File(...)):
    if len(files) != 1:
        raise HTTPException(status_code=400, detail="Please upload exactly one PDF.")
    content = await files[0].read()
    doc = fitz.open("pdf", content)
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
        for i in range(len(doc)):
            page = doc.load_page(i)
            pix = page.get_pixmap(dpi=150)
            img_bytes = pix.tobytes("jpeg")
            zip_file.writestr(f"page_{i+1}.jpg", img_bytes)
    doc.close()
    zip_buffer.seek(0)
    return Response(content=zip_buffer.read(), media_type="application/zip", headers={"Content-Disposition": "attachment; filename=pdf_images.zip"})

@app.post("/api/jpg-to-pdf")
async def jpg_to_pdf(files: List[UploadFile] = File(...)):
    if len(files) < 1:
        raise HTTPException(status_code=400, detail="Please upload at least one image.")
    images = []
    for file in files:
        if not any(file.filename.lower().endswith(ext) for ext in ['.jpg', '.jpeg', '.png']):
            raise HTTPException(status_code=400, detail="Only JPG and PNG files are allowed.")
        content = await file.read()
        try:
            img = Image.open(BytesIO(content)).convert('RGB')
            images.append(img)
        except:
            raise HTTPException(status_code=400, detail=f"Invalid image: {file.filename}")
    if not images:
        raise HTTPException(status_code=400, detail="No valid images")
    pdf_buffer = BytesIO()
    if len(images) == 1:
        images[0].save(pdf_buffer, format='PDF', save_all=True)
    else:
        images[0].save(pdf_buffer, format='PDF', save_all=True, append_images=images[1:])
    pdf_buffer.seek(0)
    return Response(content=pdf_buffer.read(), media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=images.pdf"})

@app.post("/api/rotate")
async def rotate_pdf(files: List[UploadFile] = File(...), rotation: int = Form(90)):
    if len(files) != 1:
        raise HTTPException(status_code=400, detail="Please upload exactly one PDF.")
    content = await files[0].read()
    reader = PdfReader(BytesIO(content))
    writer = PdfWriter()
    for page in reader.pages:
        page.rotate(rotation)
        writer.add_page(page)
    output = BytesIO()
    writer.write(output)
    output.seek(0)
    return Response(content=output.read(), media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=rotated.pdf"})

@app.post("/api/protect")
async def protect_pdf(files: List[UploadFile] = File(...), password: str = Form(...)):
    if len(files) != 1:
        raise HTTPException(status_code=400, detail="Please upload exactly one PDF.")
    if not password:
        raise HTTPException(status_code=400, detail="Password is required")
    content = await files[0].read()
    reader = PdfReader(BytesIO(content))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt(password)
    output = BytesIO()
    writer.write(output)
    output.seek(0)
    return Response(content=output.read(), media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=protected.pdf"})

@app.post("/api/unlock")
async def unlock_pdf(files: List[UploadFile] = File(...), password: str = Form("")):
    if len(files) != 1:
        raise HTTPException(status_code=400, detail="Please upload exactly one PDF.")
    content = await files[0].read()
    reader = PdfReader(BytesIO(content))
    if reader.is_encrypted:
        try:
            mapped = reader.decrypt(password)
            if not mapped:
                raise HTTPException(status_code=400, detail="Incorrect password or encryption not supported.")
        except Exception:
            raise HTTPException(status_code=400, detail="Incorrect password.")
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    output = BytesIO()
    writer.write(output)
    output.seek(0)
    return Response(content=output.read(), media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=unlocked.pdf"})


@app.post("/api/pdf-to-word")
async def pdf_to_word(files: List[UploadFile] = File(...)):
    if len(files) != 1:
        raise HTTPException(status_code=400, detail="Please upload exactly one PDF.")
    content = await files[0].read()
    doc = fitz.open("pdf", content)
    docx = Document()
    for i in range(len(doc)):
        page = doc.load_page(i)
        text = page.get_text()
        docx.add_paragraph(text)
        if i < len(doc) - 1:
            docx.add_page_break()
    doc.close()
    output = BytesIO()
    docx.save(output)
    output.seek(0)
    return Response(content=output.read(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": "attachment; filename=extracted.docx"})

@app.post("/api/pdf-to-txt")
async def pdf_to_txt(files: List[UploadFile] = File(...)):
    if len(files) != 1:
        raise HTTPException(status_code=400, detail="Please upload exactly one PDF.")
    content = await files[0].read()
    doc = fitz.open("pdf", content)
    text = ""
    for i in range(len(doc)):
        page = doc.load_page(i)
        text += page.get_text() + "\n--- Page Break ---\n"
    doc.close()
    return Response(content=text.encode("utf-8"), media_type="text/plain", headers={"Content-Disposition": "attachment; filename=extracted.txt"})

@app.post("/api/word-to-pdf")
async def word_to_pdf(files: List[UploadFile] = File(...)):
    if len(files) != 1:
        raise HTTPException(status_code=400, detail="Please upload exactly one DOCX.")
    if not files[0].filename.lower().endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only DOCX format supported for this simple convert.")
    content = await files[0].read()
    try:
        docx_doc = Document(BytesIO(content))
    except Exception as e:
        raise HTTPException(status_code=400, detail="Invalid DOCX file")
    
    pdf_buffer = BytesIO()
    c = canvas.Canvas(pdf_buffer, pagesize=letter)
    width, height = letter
    textobject = c.beginText()
    textobject.setTextOrigin(50, height - 50)
    textobject.setFont("Helvetica", 11)
    
    for para in docx_doc.paragraphs:
        if para.text.strip():
            lines = para.text.split('\n')
            for line in lines:
                textobject.textLine(line)
                if textobject.getY() < 50:
                    c.drawText(textobject)
                    c.showPage()
                    textobject = c.beginText()
                    textobject.setTextOrigin(50, height - 50)
                    textobject.setFont("Helvetica", 11)
        else:
            textobject.textLine("")
    
    c.drawText(textobject)
    c.save()
    pdf_buffer.seek(0)
    return Response(content=pdf_buffer.read(), media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=converted.pdf"})

@app.post("/api/watermark")
async def watermark(files: List[UploadFile] = File(...), watermark_text: str = Form("CONFIDENTIAL")):
    if len(files) != 1:
        raise HTTPException(status_code=400, detail="Please upload exactly one PDF.")
    content = await files[0].read()
    
    # Create watermark PDF
    packet = BytesIO()
    can = canvas.Canvas(packet, pagesize=letter)
    can.setFont("Helvetica-Bold", 60)
    can.setFillColorRGB(0.8, 0.8, 0.8, alpha=0.5)
    can.translate(300, 400)
    can.rotate(45)
    can.drawCentredString(0, 0, watermark_text)
    can.save()
    packet.seek(0)
    watermark_pdf = PdfReader(packet)
    watermark_page = watermark_pdf.pages[0]

    reader = PdfReader(BytesIO(content))
    writer = PdfWriter()

    for page in reader.pages:
        page.merge_page(watermark_page)
        writer.add_page(page)

    output = BytesIO()
    writer.write(output)
    output.seek(0)
    return Response(content=output.read(), media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=watermarked.pdf"})


@app.post("/api/page-numbers")
async def add_page_numbers(files: List[UploadFile] = File(...)):
    if len(files) != 1:
        raise HTTPException(status_code=400, detail="Please upload exactly one PDF.")
    content = await files[0].read()
    reader = PdfReader(BytesIO(content))
    writer = PdfWriter()

    total_pages = len(reader.pages)
    
    for i, page in enumerate(reader.pages):
        packet = BytesIO()
        can = canvas.Canvas(packet, pagesize=(page.mediabox.width, page.mediabox.height))
        can.setFont("Helvetica", 10)
        # Position at center bottom
        can.drawCentredString(float(page.mediabox.width)/2, 20, f"Page {i+1} of {total_pages}")
        can.save()
        packet.seek(0)
        num_pdf = PdfReader(packet)
        page.merge_page(num_pdf.pages[0])
        writer.add_page(page)

    output = BytesIO()
    writer.write(output)
    output.seek(0)
    return Response(content=output.read(), media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=numbered.pdf"})


@app.post("/api/delete-pages")
async def delete_pages(files: List[UploadFile] = File(...), pages_to_delete: str = Form(...)): # comma separated, 1-indexed
    if len(files) != 1:
        raise HTTPException(status_code=400, detail="Please upload exactly one PDF.")
    content = await files[0].read()
    reader = PdfReader(BytesIO(content))
    writer = PdfWriter()

    try:
        to_del = [int(p.strip()) for p in pages_to_delete.split(",")]
    except:
         raise HTTPException(status_code=400, detail="Invalid format for pages. Use comma separated numbers like 1,3,4")

    for i, page in enumerate(reader.pages):
        if (i + 1) not in to_del:
            writer.add_page(page)

    if len(writer.pages) == 0:
        raise HTTPException(status_code=400, detail="Cannot delete all pages from the PDF.")

    output = BytesIO()
    writer.write(output)
    output.seek(0)
    return Response(content=output.read(), media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=modified.pdf"})

@app.post("/api/extract-pages")
async def extract_pages(files: List[UploadFile] = File(...), pages_to_extract: str = Form(...)): # comma separated, 1-indexed
    if len(files) != 1:
        raise HTTPException(status_code=400, detail="Please upload exactly one PDF.")
    content = await files[0].read()
    reader = PdfReader(BytesIO(content))
    writer = PdfWriter()

    try:
        to_ext = [int(p.strip()) for p in pages_to_extract.split(",")]
    except:
         raise HTTPException(status_code=400, detail="Invalid format for pages. Use comma separated numbers like 1,3,4")

    for i, page in enumerate(reader.pages):
        if (i + 1) in to_ext:
            writer.add_page(page)

    if len(writer.pages) == 0:
        raise HTTPException(status_code=400, detail="No valid pages extracted. Check page numbers.")

    output = BytesIO()
    writer.write(output)
    output.seek(0)
    return Response(content=output.read(), media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=extracted.pdf"})


@app.post("/api/grayscale")
async def grayscale_pdf(files: List[UploadFile] = File(...)):
    if len(files) != 1:
        raise HTTPException(status_code=400, detail="Please upload exactly one PDF.")
    content = await files[0].read()
    
    # We can use PyMuPDF to convert each page to grayscale image and save as PDF. 
    # This might lose text selectability but effectively grayscales it.
    doc = fitz.open("pdf", content)
    pdf_buffer = BytesIO()
    images = []

    for i in range(len(doc)):
        page = doc.load_page(i)
        pix = page.get_pixmap(dpi=150, colorspace=fitz.csGRAY)
        img_bytes = pix.tobytes("jpeg")
        img = Image.open(BytesIO(img_bytes))
        images.append(img)
    
    if images:
        if len(images) == 1:
            images[0].save(pdf_buffer, format='PDF', save_all=True)
        else:
            images[0].save(pdf_buffer, format='PDF', save_all=True, append_images=images[1:])
    
    doc.close()
    pdf_buffer.seek(0)
    return Response(content=pdf_buffer.read(), media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=grayscale.pdf"})

